import glob
import docx
import os
from docx.shared import Inches

codeFilePath = '../Files/Codes/'
imageFilePath = '../Files/Images/'
docpath = '../Documents/'
if (not os.path.isdir(docpath)):
    os.mkdir(docpath)

docName = input('Enter a name for your document:')

doc = docx.Document()
docWidth = doc.sections[0].page_width

heading = input('Enter a heading for your document:')
hSize = 0

doc.add_heading(heading, hSize)
doc.paragraphs[-1].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

codeExt = input('Enter the extention for your codes:')
codeFiles = glob.glob(codeFilePath+"*.{0}".format(codeExt))
codeFont = 'Courier New'

imgExt = input('Enter the extention for your images:')
imageFiles = (glob.glob(imageFilePath+"*.{0}".format(imgExt)))


for fileIdx in range(len(codeFiles)):
    subHeading = input('Enter a subheading:')
    sub = doc.add_heading(subHeading, 2)

    cpath = codeFilePath+"{0}.{1}".format(fileIdx+1, codeExt)
    code = open(cpath)
    para = doc.add_paragraph()
    for line in code:
        temp = para.add_run(line)
        temp.font.name = codeFont
        temp.bold = True
    code.close()
    ipath = imageFilePath+"{0}.{1}".format(fileIdx+1, imgExt)
    doc.add_picture(ipath, width=Inches(docWidth.inches-2))
    doc.paragraphs[-1].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

doc.save(docpath+'{0}.docx'.format(docName))
